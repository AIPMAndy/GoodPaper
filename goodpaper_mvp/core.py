from __future__ import annotations

import json
import re
import shutil
from copy import deepcopy
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PKG_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"

NS = {
    "w": W_NS,
    "r": R_NS,
    "rel": PKG_REL_NS,
    "ct": CT_NS,
}

NUMBERING_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering"
)
NUMBERING_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml"
)

SEMANTIC_STYLE_IDS = {
    "title": "papertitle",
    "author": "author",
    "address": "address",
    "abstract": "abstract",
    "keywords": "keywords",
    "heading_1": "heading1",
    "heading_2": "heading2",
    "first_paragraph": "p1a",
    "table_caption": "tablecaption",
    "figure_caption": "figurecaption",
    "equation": "equation",
    "references": "referenceitem",
}

REQUIRED_SEMANTIC_KEYS = {
    "title",
    "author",
    "abstract",
    "keywords",
    "heading_1",
    "heading_2",
    "references",
}

BUILTIN_STYLE_IDS = {
    "Normal",
    "Header",
    "Footer",
    "FootnoteText",
    "DefaultParagraphFont",
    "Hyperlink",
    "TableNormal",
    "NoList",
}

TABLE_CAPTION_RE = re.compile(r"^table\s+\d+[\s.:]", re.IGNORECASE)
FIGURE_CAPTION_RE = re.compile(r"^(figure|fig\.)\s+\d+[\s.:]", re.IGNORECASE)
REFERENCE_ENTRY_RE = re.compile(r"^(?:\[\d+\]|\d+[.)])\s+\S")
REFERENCE_ENTRY_NUMBER_RE = re.compile(r"^(?:\[(\d+)\]|(\d+)[.)])\s+\S")
CITATION_BLOCK_RE = re.compile(r"\[(\d+(?:\s*[-,;–]\s*\d+)*)\]")
EQUATION_RE = re.compile(r".*(?:=|≤|≥|≠|\+|-|/|\*|\^).*(?:\(\d+\)|\[eq\])?$")
HEADING_TRAILING_PUNCT_RE = re.compile(r"[.!?;:。！？；：]\s*$")
KEYWORDS_PREFIXES = ("keywords", "keyword", "key words", "index terms")
UNNUMBERED_HEADING_TITLES = {
    "references",
    "reference",
    "acknowledgments",
    "acknowledgements",
    "appendix",
    "appendices",
    "conclusion",
}

HEADING_MAX_WORDS = 14
MAX_PRIMARY_SECTION_NUMBER = 20
REFERENCE_TAIL_WINDOW_RATIO = 0.55

SEVERITY_LABELS_ZH = {
    "error": "错误",
    "warn": "警告",
    "info": "提示",
}

ISSUE_CATALOG = {
    "missing_style_definition": {
        "title_zh": "文档未携带模板样式",
        "why_zh": "稿件里缺少模板定义的段落样式，系统只能按内容规则推断格式。",
        "fix_zh": "先运行一次自动排版，或让作者基于出版社模板重新另存为 .docx。",
    },
    "missing_title": {
        "title_zh": "未识别到论文标题",
        "why_zh": "文档开头没有检测到有效标题段落，后续作者信息和摘要定位会受影响。",
        "fix_zh": "确认标题位于正文最前面，并单独占一段。",
    },
    "title_style_mismatch": {
        "title_zh": "标题样式不符合模板",
        "why_zh": "标题段落存在，但没有使用模板规定的标题样式。",
        "fix_zh": "将论文标题应用为模板标题样式，或直接使用自动修复。",
    },
    "missing_abstract": {
        "title_zh": "缺少摘要段落",
        "why_zh": "系统没有找到以 'Abstract' 开头的摘要段落。",
        "fix_zh": "补充摘要，并让摘要段以 'Abstract' 开头。",
    },
    "abstract_style_mismatch": {
        "title_zh": "摘要样式不符合模板",
        "why_zh": "摘要内容存在，但段落样式不是模板要求的摘要样式。",
        "fix_zh": "把摘要段改成模板摘要样式，或使用自动修复。",
    },
    "missing_keywords": {
        "title_zh": "缺少关键词段落",
        "why_zh": "系统没有找到以 'Keywords' 开头的关键词段落。",
        "fix_zh": "补充关键词段，并以 'Keywords' 开头。",
    },
    "keywords_style_mismatch": {
        "title_zh": "关键词样式不符合模板",
        "why_zh": "关键词内容存在，但没有使用模板关键词样式。",
        "fix_zh": "把关键词段改成模板关键词样式，或使用自动修复。",
    },
    "author_style_mismatch": {
        "title_zh": "作者行样式不符合模板",
        "why_zh": "标题后的作者姓名行没有使用模板作者样式。",
        "fix_zh": "将作者姓名段应用为模板作者样式。",
    },
    "address_style_mismatch": {
        "title_zh": "单位或邮箱行样式不符合模板",
        "why_zh": "作者单位、地址或邮箱行没有使用模板地址样式。",
        "fix_zh": "将单位、地址、邮箱段应用为模板地址样式。",
    },
    "missing_references_heading": {
        "title_zh": "缺少参考文献标题",
        "why_zh": "文中已经出现参考文献条目，但没有检测到 'References' 标题。",
        "fix_zh": "在参考文献列表前补一个单独的 'References' 标题。",
    },
    "heading1_style_mismatch": {
        "title_zh": "一级标题样式不符合模板",
        "why_zh": "检测到一级标题候选段落，但样式不是模板一级标题样式。",
        "fix_zh": "将该段设置为模板一级标题样式。",
    },
    "heading2_style_mismatch": {
        "title_zh": "二级标题样式不符合模板",
        "why_zh": "检测到二级标题候选段落，但样式不是模板二级标题样式。",
        "fix_zh": "将该段设置为模板二级标题样式。",
    },
    "table_caption_style_mismatch": {
        "title_zh": "表题样式不符合模板",
        "why_zh": "表格标题存在，但没有使用模板表题样式。",
        "fix_zh": "将表题段设置为模板表题样式。",
    },
    "figure_caption_style_mismatch": {
        "title_zh": "图题样式不符合模板",
        "why_zh": "图片标题存在，但没有使用模板图题样式。",
        "fix_zh": "将图题段设置为模板图题样式。",
    },
    "equation_style_mismatch": {
        "title_zh": "公式段样式不符合模板",
        "why_zh": "检测到独立公式段，但没有使用模板公式样式。",
        "fix_zh": "将独立公式段设置为模板公式样式。",
    },
    "reference_style_mismatch": {
        "title_zh": "参考文献条目样式不符合模板",
        "why_zh": "参考文献条目没有使用模板参考文献样式。",
        "fix_zh": "将参考文献条目设置为模板参考文献样式，或使用自动修复。",
    },
    "reference_numbering_start_mismatch": {
        "title_zh": "参考文献编号不是从 [1] 开始",
        "why_zh": "参考文献列表首条编号异常，通常意味着条目缺失或编号改乱。",
        "fix_zh": "核对参考文献列表，确保第一条从 [1] 开始编号。",
    },
    "reference_numbering_gap": {
        "title_zh": "参考文献编号不连续",
        "why_zh": "参考文献编号出现跳号，说明中间有条目缺失或编号错误。",
        "fix_zh": "按顺序重排参考文献编号，确保连续递增。",
    },
    "reference_numbering_duplicate": {
        "title_zh": "参考文献编号重复",
        "why_zh": "两条参考文献使用了同一个编号。",
        "fix_zh": "检查重复条目，保证每条参考文献编号唯一。",
    },
    "reference_numbering_out_of_order": {
        "title_zh": "参考文献编号顺序倒退",
        "why_zh": "参考文献编号没有按从小到大的顺序排列。",
        "fix_zh": "重新排序或重编编号，保证参考文献按升序排列。",
    },
    "citation_missing_reference": {
        "title_zh": "正文引用没有对应参考文献条目",
        "why_zh": "正文里出现了引用编号，但参考文献列表中没有对应编号。",
        "fix_zh": "补充对应参考文献条目，或把正文引用编号改成现有参考文献编号。",
    },
    "uncited_reference_entry": {
        "title_zh": "参考文献条目未在正文中被引用",
        "why_zh": "参考文献列表中存在某条文献，但正文没有检测到对应引用编号。",
        "fix_zh": "确认该文献是否真的在正文中被引用；如未引用，建议删除或补正文引用。",
    },
    "first_paragraph_style_mismatch": {
        "title_zh": "标题后首段样式不符合模板",
        "why_zh": "一级或二级标题后的首个正文段没有使用模板首段样式。",
        "fix_zh": "将标题后的第一段正文设置为模板首段样式。",
    },
    "non_template_style_used": {
        "title_zh": "文档使用了模板外样式",
        "why_zh": "稿件中存在模板没有定义的段落样式，可能来自作者自行修改。",
        "fix_zh": "改回模板已有样式，避免混入自定义样式。",
    },
}


@dataclass
class Issue:
    severity: str
    code: str
    message: str
    paragraph_index: int | None = None
    paragraph_text: str | None = None
    expected_style: str | None = None
    actual_style: str | None = None

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


def qn(namespace: str, tag: str) -> str:
    return f"{{{namespace}}}{tag}"


def make_xml_parser() -> etree.XMLParser:
    return etree.XMLParser(remove_blank_text=True)


def severity_label_zh(severity: str) -> str:
    return SEVERITY_LABELS_ZH.get(severity, severity)


def issue_meta(code: str) -> dict[str, str]:
    return ISSUE_CATALOG.get(
        code,
        {
            "title_zh": code,
            "why_zh": "系统检测到一个未归类的问题，请结合上下文核对格式。",
            "fix_zh": "请根据模板要求手动修正后再复查。",
        },
    )


def enrich_issue_payload(payload: dict[str, Any], code: str, severity: str) -> dict[str, Any]:
    meta = issue_meta(code)
    return {
        **payload,
        "severity_label_zh": severity_label_zh(severity),
        "title_zh": meta["title_zh"],
        "why_zh": meta["why_zh"],
        "fix_zh": meta["fix_zh"],
    }


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").replace("\xa0", " ")).strip()


def normalized_lower(text: str) -> str:
    return normalize_text(text).strip(" :.-").lower()


def is_abstract_paragraph(text: str) -> bool:
    return normalized_lower(text).startswith("abstract")


def is_keywords_paragraph(text: str) -> bool:
    lowered = normalized_lower(text)
    return any(lowered.startswith(prefix) for prefix in KEYWORDS_PREFIXES)


def is_reference_heading(text: str) -> bool:
    return normalized_lower(text) in {"references", "reference"}


def is_unnumbered_heading(text: str) -> bool:
    lowered = normalized_lower(text)
    return lowered in UNNUMBERED_HEADING_TITLES


def is_table_caption(text: str) -> bool:
    return bool(TABLE_CAPTION_RE.match(normalize_text(text)))


def is_figure_caption(text: str) -> bool:
    return bool(FIGURE_CAPTION_RE.match(normalize_text(text)))


def is_reference_entry(text: str) -> bool:
    return bool(REFERENCE_ENTRY_RE.match(normalize_text(text)))


def extract_reference_number(text: str) -> int | None:
    match = REFERENCE_ENTRY_NUMBER_RE.match(normalize_text(text))
    if not match:
        return None
    return int(match.group(1) or match.group(2))


def extract_citation_numbers(text: str) -> list[int]:
    normalized = normalize_text(text)
    numbers: list[int] = []
    seen: set[int] = set()

    for match in CITATION_BLOCK_RE.finditer(normalized):
        content = match.group(1)
        for part in re.split(r"\s*[;,]\s*", content):
            token = part.strip()
            if not token:
                continue

            range_match = re.fullmatch(r"(\d+)\s*[-–]\s*(\d+)", token)
            values: list[int] = []
            if range_match:
                start = int(range_match.group(1))
                end = int(range_match.group(2))
                if start <= end:
                    values = list(range(start, end + 1))
                else:
                    values = [start, end]
            elif token.isdigit():
                values = [int(token)]

            for value in values:
                if value in seen:
                    continue
                seen.add(value)
                numbers.append(value)

    return numbers


def is_equation_like(text: str) -> bool:
    normalized = normalize_text(text)
    if not normalized or len(normalized.split()) > 18:
        return False
    if normalized.endswith("."):
        return False
    if not EQUATION_RE.match(normalized):
        return False
    if any(token in normalized for token in ("=", "≤", "≥", "≠")):
        return True
    if re.search(r"\(\d+\)\s*$", normalized) and any(
        token in normalized for token in ("+", "/", "*", "^")
    ):
        return True
    return False


def is_address_like(text: str) -> bool:
    lowered = normalized_lower(text)
    if "@" in text:
        return True
    tokens = (
        "university",
        "institute",
        "department",
        "school",
        "college",
        "laboratory",
        "lab",
        "centre",
        "center",
        "academy",
        "hospital",
        "faculty",
        "e-mail",
        "email",
    )
    return any(token in lowered for token in tokens)


def is_special_paragraph(text: str) -> bool:
    return any(
        (
            heading_level(text) is not None,
            is_unnumbered_heading(text),
            is_abstract_paragraph(text),
            is_keywords_paragraph(text),
            is_table_caption(text),
            is_figure_caption(text),
            is_reference_entry(text),
            is_equation_like(text),
        )
    )


def next_non_empty_paragraph(
    paragraphs: list[dict[str, Any]], current_index: int
) -> dict[str, Any] | None:
    for paragraph in paragraphs:
        if paragraph["index"] <= current_index:
            continue
        if paragraph["normalized_text"]:
            return paragraph
    return None


def collect_reference_entries(
    paragraphs: list[dict[str, Any]],
    references_heading_index: int | None,
) -> list[dict[str, Any]]:
    if references_heading_index is None:
        if not paragraphs:
            return []

        tail_start_index = max(1, int(len(paragraphs) * REFERENCE_TAIL_WINDOW_RATIO))
        fallback_entries = [
            paragraph
            for paragraph in paragraphs
            if paragraph["index"] >= tail_start_index and is_reference_entry(paragraph["text"])
        ]
        if not fallback_entries:
            return []

        first_entry_index = fallback_entries[0]["index"]
        entries: list[dict[str, Any]] = []
        block_started = False
        for paragraph in paragraphs:
            if paragraph["index"] < first_entry_index:
                continue
            text = paragraph["normalized_text"]
            if not text:
                continue
            if is_reference_entry(text):
                entries.append(paragraph)
                block_started = True
                continue
            if block_started:
                break

        return entries or fallback_entries

    entries: list[dict[str, Any]] = []
    block_started = False
    for paragraph in paragraphs:
        if paragraph["index"] <= references_heading_index:
            continue
        text = paragraph["normalized_text"]
        if not text:
            continue
        if is_reference_entry(text):
            entries.append(paragraph)
            block_started = True
            continue
        if block_started:
            break

    return entries or fallback_entries


def collect_citation_occurrences(
    paragraphs: list[dict[str, Any]],
    references_heading_index: int | None,
    reference_entries: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    occurrences: list[dict[str, Any]] = []
    reference_entry_indices = {paragraph["index"] for paragraph in reference_entries}

    for paragraph in paragraphs:
        if references_heading_index is not None and paragraph["index"] >= references_heading_index:
            continue
        if paragraph["index"] in reference_entry_indices:
            continue
        text = paragraph["normalized_text"]
        if not text or is_reference_entry(text):
            continue
        numbers = extract_citation_numbers(text)
        for number in numbers:
            occurrences.append(
                {
                    "paragraph_index": paragraph["index"],
                    "paragraph_text": text,
                    "style_id": paragraph.get("style_id"),
                    "citation_number": number,
                }
            )

    return occurrences


def recommendations_from_issues(issues: list[Issue]) -> list[str]:
    codes = {issue.code for issue in issues}
    recommendations: list[str] = []

    if any(code.endswith("style_mismatch") or code == "missing_style_definition" for code in codes):
        recommendations.append("先运行一次自动排版，再重新执行格式检查。")
    if "missing_abstract" in codes:
        recommendations.append("补充摘要段落，并确保以 'Abstract' 开头。")
    if "missing_keywords" in codes:
        recommendations.append("补充关键词段落，并确保以 'Keywords' 开头。")
    if "missing_references_heading" in codes:
        recommendations.append("为参考文献列表补一个 'References' 标题。")
    if any(code in codes for code in {"table_caption_style_mismatch", "figure_caption_style_mismatch"}):
        recommendations.append("检查图表标题是否分别使用表题和图题样式。")
    if "equation_style_mismatch" in codes:
        recommendations.append("检查独立公式段落是否使用模板里的公式样式。")
    if any(
        code in codes
        for code in {
            "reference_numbering_start_mismatch",
            "reference_numbering_gap",
            "reference_numbering_duplicate",
            "reference_numbering_out_of_order",
        }
    ):
        recommendations.append("核对参考文献编号是否从 [1] 开始且连续递增，并与正文引用保持一致。")
    if "citation_missing_reference" in codes:
        recommendations.append("检查正文引用编号是否都能在参考文献列表中找到对应条目。")
    if "uncited_reference_entry" in codes:
        recommendations.append("检查参考文献列表是否存在未被正文引用的冗余条目。")

    return recommendations


def summarize_codes(issues: list[Issue]) -> list[dict[str, Any]]:
    severity_weight = {"error": 3, "warn": 2, "info": 1}
    counts: dict[str, dict[str, Any]] = {}
    for issue in issues:
        meta = issue_meta(issue.code)
        entry = counts.setdefault(
            issue.code,
            {
                "code": issue.code,
                "severity": issue.severity,
                "severity_label_zh": severity_label_zh(issue.severity),
                "count": 0,
                "message": issue.message,
                "title_zh": meta["title_zh"],
                "why_zh": meta["why_zh"],
                "fix_zh": meta["fix_zh"],
            },
        )
        entry["count"] += 1
        if severity_weight[issue.severity] > severity_weight[entry["severity"]]:
            entry["severity"] = issue.severity
            entry["severity_label_zh"] = severity_label_zh(issue.severity)

    return sorted(
        counts.values(),
        key=lambda item: (severity_weight[item["severity"]], item["count"]),
        reverse=True,
    )


def build_paragraph_findings(
    paragraphs: list[dict[str, Any]],
    issues: list[Issue],
) -> list[dict[str, Any]]:
    by_index = {paragraph["index"]: paragraph for paragraph in paragraphs}
    grouped: dict[int, dict[str, Any]] = {}

    for issue in issues:
        if issue.paragraph_index is None:
            continue
        paragraph = by_index.get(issue.paragraph_index)
        if paragraph is None:
            continue
        entry = grouped.setdefault(
            issue.paragraph_index,
            {
                "paragraph_index": issue.paragraph_index,
                "paragraph_text": paragraph["normalized_text"],
                "current_style": paragraph.get("style_id"),
                "messages": [],
                "codes": [],
                "expected_styles": [],
                "severities": [],
                "auto_fixable": False,
            },
        )
        entry["messages"].append(issue.message)
        entry["codes"].append(issue.code)
        entry["severities"].append(issue.severity)
        meta = issue_meta(issue.code)
        entry.setdefault("titles_zh", [])
        entry.setdefault("suggestions", [])
        if meta["title_zh"] not in entry["titles_zh"]:
            entry["titles_zh"].append(meta["title_zh"])
        if meta["fix_zh"] not in entry["suggestions"]:
            entry["suggestions"].append(meta["fix_zh"])
        if issue.expected_style and issue.expected_style not in entry["expected_styles"]:
            entry["expected_styles"].append(issue.expected_style)
        if issue.expected_style and paragraph.get("style_id") != issue.expected_style:
            entry["auto_fixable"] = True

    severity_rank = {"error": 3, "warn": 2, "info": 1}
    findings: list[dict[str, Any]] = []
    for entry in grouped.values():
        primary_severity = max(entry["severities"], key=lambda item: severity_rank[item])
        findings.append(
            {
                "paragraph_index": entry["paragraph_index"],
                "paragraph_text": entry["paragraph_text"],
                "current_style": entry["current_style"],
                "primary_severity": primary_severity,
                "primary_severity_label_zh": severity_label_zh(primary_severity),
                "issue_count": len(entry["codes"]),
                "codes": entry["codes"],
                "messages": entry["messages"],
                "titles_zh": entry["titles_zh"],
                "suggestions": entry["suggestions"],
                "expected_styles": entry["expected_styles"],
                "auto_fixable": entry["auto_fixable"],
            }
        )

    return sorted(findings, key=lambda item: item["paragraph_index"])


def build_fix_plan(
    paragraphs: list[dict[str, Any]],
    issues: list[Issue],
) -> dict[str, Any]:
    by_index = {paragraph["index"]: paragraph for paragraph in paragraphs}
    actions: dict[tuple[int, str], dict[str, Any]] = {}
    non_fixable_codes: list[str] = []

    for issue in issues:
        if issue.paragraph_index is None or not issue.expected_style:
            if issue.code not in non_fixable_codes:
                non_fixable_codes.append(issue.code)
            continue
        paragraph = by_index.get(issue.paragraph_index)
        if paragraph is None:
            continue
        if paragraph.get("style_id") == issue.expected_style:
            continue
        key = (issue.paragraph_index, issue.expected_style)
        entry = actions.setdefault(
            key,
            {
                "paragraph_index": issue.paragraph_index,
                "paragraph_text": paragraph["normalized_text"],
                "from_style": paragraph.get("style_id"),
                "to_style": issue.expected_style,
                "codes": [],
                "messages": [],
                "titles_zh": [],
                "suggestions": [],
            },
        )
        entry["codes"].append(issue.code)
        entry["messages"].append(issue.message)
        meta = issue_meta(issue.code)
        if meta["title_zh"] not in entry["titles_zh"]:
            entry["titles_zh"].append(meta["title_zh"])
        if meta["fix_zh"] not in entry["suggestions"]:
            entry["suggestions"].append(meta["fix_zh"])

    action_list = sorted(actions.values(), key=lambda item: item["paragraph_index"])
    return {
        "auto_fixable_issue_count": len(action_list),
        "actions": action_list,
        "non_fixable_codes": non_fixable_codes,
    }


def read_package_part(package_path: Path, part_name: str) -> bytes:
    with ZipFile(package_path) as archive:
        return archive.read(part_name)


def parse_package_xml(package_path: Path, part_name: str) -> etree._Element:
    return etree.fromstring(
        read_package_part(package_path, part_name),
        parser=make_xml_parser(),
    )


def get_style_definitions(package_path: Path) -> dict[str, dict[str, Any]]:
    styles_root = parse_package_xml(package_path, "word/styles.xml")
    style_map: dict[str, dict[str, Any]] = {}

    for style in styles_root.findall("w:style", NS):
        style_id = style.get(qn(W_NS, "styleId"))
        style_type = style.get(qn(W_NS, "type"))
        if not style_id or not style_type:
            continue
        name_el = style.find("w:name", NS)
        style_map[style_id] = {
            "style_id": style_id,
            "name": name_el.get(qn(W_NS, "val")) if name_el is not None else style_id,
            "type": style_type,
            "custom": style.get(qn(W_NS, "customStyle")) == "1",
        }

    return style_map


def get_template_profile(
    template_path: Path, semantic_style_ids: dict[str, str] | None = None
) -> dict[str, Any]:
    styles = get_style_definitions(template_path)
    paragraph_styles = {
        style_id: meta for style_id, meta in styles.items() if meta["type"] == "paragraph"
    }
    semantic_style_ids = semantic_style_ids or SEMANTIC_STYLE_IDS
    semantic_styles = {
        key: style_id
        for key, style_id in semantic_style_ids.items()
        if style_id in paragraph_styles
    }
    return {
        "template_path": str(template_path),
        "paragraph_styles": paragraph_styles,
        "semantic_styles": semantic_styles,
        "paragraph_style_ids": sorted(paragraph_styles),
        "custom_paragraph_style_ids": sorted(
            style_id
            for style_id, meta in paragraph_styles.items()
            if meta["custom"]
        ),
    }


def get_document_paragraphs(document_path: Path) -> list[dict[str, Any]]:
    document = Document(document_path)
    paragraphs: list[dict[str, Any]] = []

    for index, paragraph in enumerate(document.paragraphs, start=1):
        style_name = None
        style_id = None
        try:
            if paragraph.style is not None:
                style_name = paragraph.style.name
                style_id = paragraph.style.style_id
        except Exception:
            pass
        paragraphs.append(
            {
                "index": index,
                "text": paragraph.text or "",
                "normalized_text": normalize_text(paragraph.text or ""),
                "style_name": style_name,
                "style_id": style_id,
            }
        )

    return paragraphs


def heading_level(text: str) -> int | None:
    normalized = normalize_text(text)
    match = re.match(r"^(\d+(?:\.\d+)*)[.)]?\s+(.+)$", normalized)
    if not match:
        return None

    sections = [part for part in match.group(1).split(".") if part]
    title = match.group(2).strip()
    if not sections or not title:
        return None
    if int(sections[0]) > MAX_PRIMARY_SECTION_NUMBER:
        return None
    if len(title.split()) > HEADING_MAX_WORDS:
        return None
    if HEADING_TRAILING_PUNCT_RE.search(title):
        return None

    if len(sections) == 1:
        return 1
    if len(sections) >= 2:
        return 2
    return None


def find_first_non_empty(paragraphs: list[dict[str, Any]]) -> int | None:
    for paragraph in paragraphs:
        if paragraph["normalized_text"]:
            return paragraph["index"]
    return None


def find_first_matching(
    paragraphs: list[dict[str, Any]], predicate
) -> int | None:
    for paragraph in paragraphs:
        if predicate(paragraph):
            return paragraph["index"]
    return None


def add_issue(issues: list[Issue], **kwargs: Any) -> None:
    issues.append(Issue(**kwargs))


def add_reference_numbering_issues(
    issues: list[Issue],
    reference_entries: list[dict[str, Any]],
) -> None:
    numbered_entries = []
    for paragraph in reference_entries:
        number = extract_reference_number(paragraph["text"])
        if number is not None:
            numbered_entries.append((paragraph, number))

    if not numbered_entries:
        return

    first_paragraph, first_number = numbered_entries[0]
    if first_number != 1:
        add_issue(
            issues,
            severity="warn",
            code="reference_numbering_start_mismatch",
            message=f"The first reference entry starts at [{first_number}] instead of [1].",
            paragraph_index=first_paragraph["index"],
            paragraph_text=first_paragraph["normalized_text"],
            actual_style=first_paragraph["style_id"],
        )

    previous_number = first_number
    for paragraph, current_number in numbered_entries[1:]:
        if current_number == previous_number:
            add_issue(
                issues,
                severity="warn",
                code="reference_numbering_duplicate",
                message=f"Reference numbering repeats [{current_number}].",
                paragraph_index=paragraph["index"],
                paragraph_text=paragraph["normalized_text"],
                actual_style=paragraph["style_id"],
            )
        elif current_number < previous_number:
            add_issue(
                issues,
                severity="warn",
                code="reference_numbering_out_of_order",
                message=(
                    f"Reference numbering goes backwards from "
                    f"[{previous_number}] to [{current_number}]."
                ),
                paragraph_index=paragraph["index"],
                paragraph_text=paragraph["normalized_text"],
                actual_style=paragraph["style_id"],
            )
        elif current_number > previous_number + 1:
            add_issue(
                issues,
                severity="warn",
                code="reference_numbering_gap",
                message=(
                    f"Reference numbering jumps from "
                    f"[{previous_number}] to [{current_number}]."
                ),
                paragraph_index=paragraph["index"],
                paragraph_text=paragraph["normalized_text"],
                actual_style=paragraph["style_id"],
            )
        previous_number = current_number


def add_citation_reference_crosscheck_issues(
    issues: list[Issue],
    citation_occurrences: list[dict[str, Any]],
    reference_entries: list[dict[str, Any]],
) -> None:
    reference_by_number: dict[int, dict[str, Any]] = {}
    for paragraph in reference_entries:
        number = extract_reference_number(paragraph["text"])
        if number is None or number in reference_by_number:
            continue
        reference_by_number[number] = paragraph

    citation_by_number: dict[int, dict[str, Any]] = {}
    for occurrence in citation_occurrences:
        number = occurrence["citation_number"]
        citation_by_number.setdefault(number, occurrence)

    for number in sorted(citation_by_number):
        if number in reference_by_number:
            continue
        occurrence = citation_by_number[number]
        add_issue(
            issues,
            severity="warn",
            code="citation_missing_reference",
            message=(
                f"Body citation [{number}] was detected, but there is no matching "
                f"reference entry."
            ),
            paragraph_index=occurrence["paragraph_index"],
            paragraph_text=occurrence["paragraph_text"],
            actual_style=occurrence["style_id"],
        )

    cited_numbers = set(citation_by_number)
    for number in sorted(reference_by_number):
        if number in cited_numbers:
            continue
        paragraph = reference_by_number[number]
        add_issue(
            issues,
            severity="info",
            code="uncited_reference_entry",
            message=(
                f"Reference entry [{number}] exists in the list, but no body citation "
                f"was detected."
            ),
            paragraph_index=paragraph["index"],
            paragraph_text=paragraph["normalized_text"],
            actual_style=paragraph["style_id"],
        )


def analyze_document(
    template_path: Path,
    document_path: Path,
    semantic_style_ids: dict[str, str] | None = None,
) -> dict[str, Any]:
    template_profile = get_template_profile(template_path, semantic_style_ids)
    document_styles = get_style_definitions(document_path)
    document_paragraph_styles = {
        style_id: meta
        for style_id, meta in document_styles.items()
        if meta["type"] == "paragraph"
    }
    paragraphs = get_document_paragraphs(document_path)
    issues: list[Issue] = []
    semantic_styles = template_profile["semantic_styles"]

    for key, style_id in semantic_styles.items():
        if key not in REQUIRED_SEMANTIC_KEYS:
            continue
        if style_id not in document_paragraph_styles:
            add_issue(
                issues,
                severity="warn",
                code="missing_style_definition",
                message=f"Document does not contain template style '{style_id}'.",
                expected_style=style_id,
            )

    title_index = find_first_non_empty(paragraphs)
    if title_index is None:
        add_issue(
            issues,
            severity="error",
            code="missing_title",
            message="No non-empty title paragraph was detected.",
        )
    elif "title" in semantic_styles:
        title_paragraph = paragraphs[title_index - 1]
        if title_paragraph["style_id"] != semantic_styles["title"]:
            add_issue(
                issues,
                severity="warn",
                code="title_style_mismatch",
                message="The first non-empty paragraph is not using the template title style.",
                paragraph_index=title_paragraph["index"],
                paragraph_text=title_paragraph["normalized_text"],
                expected_style=semantic_styles["title"],
                actual_style=title_paragraph["style_id"],
            )

    abstract_index = find_first_matching(paragraphs, lambda item: is_abstract_paragraph(item["text"]))
    if "abstract" in semantic_styles:
        if abstract_index is None:
            add_issue(
                issues,
                severity="error",
                code="missing_abstract",
                message="No paragraph starting with 'Abstract' was detected.",
                expected_style=semantic_styles["abstract"],
            )
        else:
            abstract_paragraph = paragraphs[abstract_index - 1]
            if abstract_paragraph["style_id"] != semantic_styles["abstract"]:
                add_issue(
                    issues,
                    severity="warn",
                    code="abstract_style_mismatch",
                    message="The abstract paragraph is not using the template abstract style.",
                    paragraph_index=abstract_paragraph["index"],
                    paragraph_text=abstract_paragraph["normalized_text"],
                    expected_style=semantic_styles["abstract"],
                    actual_style=abstract_paragraph["style_id"],
                )

    keywords_index = find_first_matching(paragraphs, lambda item: is_keywords_paragraph(item["text"]))
    if "keywords" in semantic_styles:
        if keywords_index is None:
            add_issue(
                issues,
                severity="warn",
                code="missing_keywords",
                message="No paragraph starting with 'Keywords' was detected.",
                expected_style=semantic_styles["keywords"],
            )
        else:
            keywords_paragraph = paragraphs[keywords_index - 1]
            if keywords_paragraph["style_id"] != semantic_styles["keywords"]:
                add_issue(
                    issues,
                    severity="warn",
                    code="keywords_style_mismatch",
                    message="The keywords paragraph is not using the template keywords style.",
                    paragraph_index=keywords_paragraph["index"],
                    paragraph_text=keywords_paragraph["normalized_text"],
                    expected_style=semantic_styles["keywords"],
                    actual_style=keywords_paragraph["style_id"],
                )

    address_candidates: list[dict[str, Any]] = []
    if title_index is not None:
        candidate_author_index = None
        boundary = abstract_index or (len(paragraphs) + 1)
        for paragraph in paragraphs:
            if paragraph["index"] <= title_index or paragraph["index"] >= boundary:
                continue
            if not paragraph["normalized_text"]:
                continue
            if candidate_author_index is None:
                candidate_author_index = paragraph["index"]
                continue
            if is_address_like(paragraph["text"]):
                address_candidates.append(paragraph)

        if candidate_author_index is not None and "author" in semantic_styles:
            author_paragraph = paragraphs[candidate_author_index - 1]
            if author_paragraph["style_id"] != semantic_styles["author"]:
                add_issue(
                    issues,
                    severity="info",
                    code="author_style_mismatch",
                    message="The first paragraph after the title is not using the template author style.",
                    paragraph_index=author_paragraph["index"],
                    paragraph_text=author_paragraph["normalized_text"],
                    expected_style=semantic_styles["author"],
                    actual_style=author_paragraph["style_id"],
                )

    if "address" in semantic_styles:
        for paragraph in address_candidates:
            if paragraph["style_id"] != semantic_styles["address"]:
                add_issue(
                    issues,
                    severity="info",
                    code="address_style_mismatch",
                    message="An affiliation or email paragraph is not using the template address style.",
                    paragraph_index=paragraph["index"],
                    paragraph_text=paragraph["normalized_text"],
                    expected_style=semantic_styles["address"],
                    actual_style=paragraph["style_id"],
                )

    references_heading_index = find_first_matching(
        paragraphs, lambda item: is_reference_heading(item["text"])
    )
    reference_entries = collect_reference_entries(paragraphs, references_heading_index)
    citation_occurrences = collect_citation_occurrences(
        paragraphs,
        references_heading_index,
        reference_entries,
    )
    if reference_entries and references_heading_index is None:
        add_issue(
            issues,
            severity="warn",
            code="missing_references_heading",
            message="Reference entries were detected, but there is no 'References' heading.",
            expected_style=semantic_styles.get("heading_1"),
        )

    for paragraph in paragraphs:
        text = paragraph["normalized_text"]
        if not text:
            continue
        level = heading_level(text)
        is_top_heading = level == 1 or is_unnumbered_heading(text)
        if is_top_heading and "heading_1" in semantic_styles:
            if paragraph["style_id"] != semantic_styles["heading_1"]:
                add_issue(
                    issues,
                    severity="info",
                    code="heading1_style_mismatch",
                    message="A level-1 heading candidate is not using the template heading1 style.",
                    paragraph_index=paragraph["index"],
                    paragraph_text=text,
                    expected_style=semantic_styles["heading_1"],
                    actual_style=paragraph["style_id"],
                )
        if level == 2 and "heading_2" in semantic_styles:
            if paragraph["style_id"] != semantic_styles["heading_2"]:
                add_issue(
                    issues,
                    severity="info",
                    code="heading2_style_mismatch",
                    message="A level-2 heading candidate is not using the template heading2 style.",
                    paragraph_index=paragraph["index"],
                    paragraph_text=text,
                    expected_style=semantic_styles["heading_2"],
                    actual_style=paragraph["style_id"],
                )
        if is_table_caption(text) and "table_caption" in semantic_styles:
            if paragraph["style_id"] != semantic_styles["table_caption"]:
                add_issue(
                    issues,
                    severity="info",
                    code="table_caption_style_mismatch",
                    message="A table caption is not using the template table caption style.",
                    paragraph_index=paragraph["index"],
                    paragraph_text=text,
                    expected_style=semantic_styles["table_caption"],
                    actual_style=paragraph["style_id"],
                )
        if is_figure_caption(text) and "figure_caption" in semantic_styles:
            if paragraph["style_id"] != semantic_styles["figure_caption"]:
                add_issue(
                    issues,
                    severity="info",
                    code="figure_caption_style_mismatch",
                    message="A figure caption is not using the template figure caption style.",
                    paragraph_index=paragraph["index"],
                    paragraph_text=text,
                    expected_style=semantic_styles["figure_caption"],
                    actual_style=paragraph["style_id"],
                )
        if is_equation_like(text) and "equation" in semantic_styles:
            if paragraph["style_id"] != semantic_styles["equation"]:
                add_issue(
                    issues,
                    severity="info",
                    code="equation_style_mismatch",
                    message="An equation-like paragraph is not using the template equation style.",
                    paragraph_index=paragraph["index"],
                    paragraph_text=text,
                    expected_style=semantic_styles["equation"],
                    actual_style=paragraph["style_id"],
                )

    if "references" in semantic_styles:
        for paragraph in reference_entries:
            if paragraph["style_id"] != semantic_styles["references"]:
                add_issue(
                    issues,
                    severity="info",
                    code="reference_style_mismatch",
                    message="A paragraph in the references block is not using the template reference style.",
                    paragraph_index=paragraph["index"],
                    paragraph_text=paragraph["normalized_text"],
                    expected_style=semantic_styles["references"],
                    actual_style=paragraph["style_id"],
                )

    add_reference_numbering_issues(issues, reference_entries)
    add_citation_reference_crosscheck_issues(
        issues,
        citation_occurrences,
        reference_entries,
    )

    if "first_paragraph" in semantic_styles:
        for paragraph in paragraphs:
            text = paragraph["normalized_text"]
            if not text:
                continue
            if heading_level(text) not in {1, 2} and not is_unnumbered_heading(text):
                continue
            next_paragraph = next_non_empty_paragraph(paragraphs, paragraph["index"])
            if next_paragraph is None or is_special_paragraph(next_paragraph["text"]):
                continue
            if next_paragraph["style_id"] != semantic_styles["first_paragraph"]:
                add_issue(
                    issues,
                    severity="info",
                    code="first_paragraph_style_mismatch",
                    message="The first body paragraph after a heading is not using the template first-paragraph style.",
                    paragraph_index=next_paragraph["index"],
                    paragraph_text=next_paragraph["normalized_text"],
                    expected_style=semantic_styles["first_paragraph"],
                    actual_style=next_paragraph["style_id"],
                )

    template_style_ids = set(template_profile["paragraph_style_ids"])
    used_style_ids = sorted(
        {
            paragraph["style_id"]
            for paragraph in paragraphs
            if paragraph["style_id"] and paragraph["style_id"] not in BUILTIN_STYLE_IDS
        }
    )
    extra_style_ids = [
        style_id for style_id in used_style_ids if style_id not in template_style_ids
    ]
    for style_id in extra_style_ids:
        add_issue(
            issues,
            severity="info",
            code="non_template_style_used",
            message=f"Document uses paragraph style '{style_id}' which is not defined in the template profile.",
            actual_style=style_id,
        )

    severity_weight = {"error": 3, "warn": 2, "info": 1}
    issues.sort(key=lambda item: severity_weight[item.severity], reverse=True)

    style_usage: dict[str, int] = {}
    for paragraph in paragraphs:
        style_id = paragraph["style_id"] or "(no-style)"
        style_usage[style_id] = style_usage.get(style_id, 0) + 1

    return {
        "template": {
            "path": str(template_path),
            "semantic_styles": semantic_styles,
            "paragraph_style_count": len(template_profile["paragraph_style_ids"]),
        },
        "document": {
            "file_name": document_path.name,
            "path": str(document_path),
            "paragraph_count": len(paragraphs),
            "paragraph_style_count": len(document_paragraph_styles),
            "style_usage": style_usage,
        },
        "issues": [
            enrich_issue_payload(issue.to_dict(), issue.code, issue.severity)
            for issue in issues
        ],
        "issue_code_summary": summarize_codes(issues),
        "paragraph_findings": build_paragraph_findings(paragraphs, issues),
        "summary": {
            "error_count": sum(issue.severity == "error" for issue in issues),
            "warn_count": sum(issue.severity == "warn" for issue in issues),
            "info_count": sum(issue.severity == "info" for issue in issues),
        },
        "fix_plan": build_fix_plan(paragraphs, issues),
        "recommendations": recommendations_from_issues(issues),
    }


def merge_styles_xml(source_bytes: bytes, template_bytes: bytes) -> bytes:
    source_root = etree.fromstring(source_bytes, parser=make_xml_parser())
    template_root = etree.fromstring(template_bytes, parser=make_xml_parser())

    source_index = {
        style.get(qn(W_NS, "styleId")): style
        for style in source_root.findall("w:style", NS)
        if style.get(qn(W_NS, "styleId"))
    }

    for template_style in template_root.findall("w:style", NS):
        style_id = template_style.get(qn(W_NS, "styleId"))
        if not style_id:
            continue
        replacement = deepcopy(template_style)
        existing = source_index.get(style_id)
        if existing is None:
            source_root.append(replacement)
        else:
            existing.getparent().replace(existing, replacement)

    return etree.tostring(
        source_root, encoding="UTF-8", xml_declaration=True, standalone="yes"
    )


def ensure_numbering_relationship(rels_bytes: bytes) -> bytes:
    root = etree.fromstring(rels_bytes, parser=make_xml_parser())
    relationships = root.findall(f"{{{PKG_REL_NS}}}Relationship")
    for relationship in relationships:
        if relationship.get("Type") == NUMBERING_REL_TYPE:
            return etree.tostring(root, encoding="UTF-8", xml_declaration=True)

    next_index = 1
    for relationship in relationships:
        rel_id = relationship.get("Id", "")
        match = re.fullmatch(r"rId(\d+)", rel_id)
        if match:
            next_index = max(next_index, int(match.group(1)) + 1)

    relationship = etree.SubElement(root, qn(PKG_REL_NS, "Relationship"))
    relationship.set("Id", f"rId{next_index}")
    relationship.set("Type", NUMBERING_REL_TYPE)
    relationship.set("Target", "numbering.xml")
    return etree.tostring(root, encoding="UTF-8", xml_declaration=True)


def ensure_numbering_override(content_types_bytes: bytes) -> bytes:
    root = etree.fromstring(content_types_bytes, parser=make_xml_parser())
    for override in root.findall("ct:Override", NS):
        if override.get("PartName") == "/word/numbering.xml":
            return etree.tostring(root, encoding="UTF-8", xml_declaration=True)

    override = etree.SubElement(root, qn(CT_NS, "Override"))
    override.set("PartName", "/word/numbering.xml")
    override.set("ContentType", NUMBERING_CONTENT_TYPE)
    return etree.tostring(root, encoding="UTF-8", xml_declaration=True)


def read_zip_entries(package_path: Path) -> dict[str, bytes]:
    with ZipFile(package_path) as archive:
        return {name: archive.read(name) for name in archive.namelist()}


def write_zip_entries(package_path: Path, entries: dict[str, bytes]) -> None:
    with ZipFile(package_path, "w", ZIP_DEFLATED) as archive:
        for name, content in entries.items():
            archive.writestr(name, content)


def merge_template_parts(template_path: Path, output_path: Path) -> None:
    template_entries = read_zip_entries(template_path)
    output_entries = read_zip_entries(output_path)

    if "word/styles.xml" in output_entries and "word/styles.xml" in template_entries:
        output_entries["word/styles.xml"] = merge_styles_xml(
            output_entries["word/styles.xml"],
            template_entries["word/styles.xml"],
        )

    if "word/numbering.xml" in template_entries:
        output_entries["word/numbering.xml"] = template_entries["word/numbering.xml"]
        output_entries["[Content_Types].xml"] = ensure_numbering_override(
            output_entries["[Content_Types].xml"]
        )
        rels_name = "word/_rels/document.xml.rels"
        if rels_name in output_entries:
            output_entries[rels_name] = ensure_numbering_relationship(output_entries[rels_name])

    write_zip_entries(output_path, output_entries)


def apply_style(paragraph, style_name: str, style_changes: list[dict[str, Any]]) -> None:
    current_style = getattr(paragraph.style, "name", None)
    if current_style == style_name:
        return
    paragraph.style = style_name
    style_changes.append(
        {
            "paragraph_text": normalize_text(paragraph.text),
            "from": current_style,
            "to": style_name,
        }
    )


def format_document(
    template_path: Path,
    input_path: Path,
    output_path: Path,
    semantic_style_ids: dict[str, str] | None = None,
) -> dict[str, Any]:
    if input_path.suffix.lower() != ".docx":
        raise ValueError("Formatting currently supports .docx manuscripts on Mac.")

    shutil.copy2(input_path, output_path)
    merge_template_parts(template_path, output_path)

    document = Document(output_path)
    template_profile = get_template_profile(template_path, semantic_style_ids)
    paragraph_styles = template_profile["paragraph_styles"]
    semantic_styles = template_profile["semantic_styles"]
    available_style_names = {style.name for style in document.styles}
    style_name_by_id = {
        style_id: meta["name"]
        for style_id, meta in paragraph_styles.items()
        if meta["name"] in available_style_names
    }

    style_changes: list[dict[str, Any]] = []
    paragraphs = list(document.paragraphs)
    paragraph_payloads = [
        {
            "index": idx + 1,
            "text": paragraph.text or "",
            "normalized_text": normalize_text(paragraph.text or ""),
        }
        for idx, paragraph in enumerate(paragraphs)
    ]

    def style_exists(style_id: str) -> bool:
        return style_id in style_name_by_id

    def set_style_if_available(index: int | None, semantic_key: str) -> None:
        style_id = semantic_styles.get(semantic_key)
        if index is None or style_id is None or not style_exists(style_id):
            return
        paragraph = paragraphs[index - 1]
        apply_style(paragraph, style_name_by_id[style_id], style_changes)

    non_empty_indices = [
        paragraph["index"] for paragraph in paragraph_payloads if paragraph["normalized_text"]
    ]
    title_index = non_empty_indices[0] if non_empty_indices else None
    abstract_index = find_first_matching(
        paragraph_payloads,
        lambda item: is_abstract_paragraph(item["text"]),
    )
    keywords_index = find_first_matching(
        paragraph_payloads,
        lambda item: is_keywords_paragraph(item["text"]),
    )

    set_style_if_available(title_index, "title")

    if title_index is not None:
        boundary = abstract_index or (len(paragraphs) + 1)
        author_indices: list[int] = []
        address_indices: list[int] = []
        for payload in paragraph_payloads:
            idx = payload["index"]
            if not (title_index < idx < boundary) or not payload["normalized_text"]:
                continue
            if is_address_like(payload["text"]):
                address_indices.append(idx)
            else:
                author_indices.append(idx)

        for idx in author_indices:
            set_style_if_available(idx, "author")
        for idx in address_indices:
            set_style_if_available(idx, "address")

    set_style_if_available(abstract_index, "abstract")
    set_style_if_available(keywords_index, "keywords")

    references_heading_index = None
    for idx, payload in enumerate(paragraph_payloads, start=1):
        text = payload["normalized_text"]
        if not text:
            continue
        if heading_level(text) == 1 or is_unnumbered_heading(text):
            set_style_if_available(idx, "heading_1")
        elif heading_level(text) == 2:
            set_style_if_available(idx, "heading_2")
        if is_table_caption(text):
            set_style_if_available(idx, "table_caption")
        if is_figure_caption(text):
            set_style_if_available(idx, "figure_caption")
        if is_equation_like(text):
            set_style_if_available(idx, "equation")
        if is_reference_heading(text):
            references_heading_index = idx
            set_style_if_available(idx, "heading_1")
    reference_entries = collect_reference_entries(paragraph_payloads, references_heading_index)
    for paragraph in reference_entries:
        set_style_if_available(paragraph["index"], "references")

    for payload in paragraph_payloads:
        text = payload["normalized_text"]
        if heading_level(text) not in {1, 2} and not is_unnumbered_heading(text):
            continue
        next_paragraph = next_non_empty_paragraph(paragraph_payloads, payload["index"])
        if next_paragraph is None or is_special_paragraph(next_paragraph["text"]):
            continue
        set_style_if_available(next_paragraph["index"], "first_paragraph")

    document.save(output_path)

    applied_style_counts: dict[str, int] = {}
    for change in style_changes:
        style_name = change["to"]
        applied_style_counts[style_name] = applied_style_counts.get(style_name, 0) + 1

    return {
        "output_path": str(output_path),
        "change_count": len(style_changes),
        "applied_style_counts": applied_style_counts,
        "style_changes": style_changes[:25],
    }


def to_pretty_json(payload: dict[str, Any]) -> str:
    return json.dumps(payload, ensure_ascii=False, indent=2)
